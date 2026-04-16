import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(doc):
    """Add a visual separator line between sections."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_label_badge(doc, text: str, color_hex: str = "4472C4"):
    """Add a colored label badge."""
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)
    return p

def add_table_to_doc(doc, structured_data):
    """Render structured table data (rows/headers) as a Word table."""
    if not structured_data:
        return

    headers = structured_data.get("headers", [])
    rows    = structured_data.get("rows", [])

    if not headers and not rows:
        return

    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'

    # Header Row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        if i < col_count:
            cell = hdr_row.cells[i]
            cell.text = str(h)
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '2E74B5')
            cell._tc.get_or_add_tcPr().append(shd)

    # Data Rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = 'EBF3FB' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            if c_idx < col_count:
                cell = row.cells[c_idx]
                cell.text = str(val)
                # Ensure at least one run exists
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill)
                cell._tc.get_or_add_tcPr().append(shd)

def generate_word_report(json_data, output_path: str):
    """Generate Word document from extraction JSON results."""
    doc = Document()
    pdf_name = Path(json_data.get("pdf", "document")).stem

    # Title
    title = doc.add_heading(f"PDF Extraction Report: {pdf_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Technical Summary
    doc.add_heading("Technical Intelligence Summary", level=1)
    requests = json_data.get("request_count", 0)
    tokens = json_data.get("token_usage", 0)
    
    p = doc.add_paragraph()
    p.add_run("• Total AI Model Interactions: ").bold = True
    p.add_run(f"{requests} requests")
    
    p = doc.add_paragraph()
    p.add_run("• Total Token Consumption: ").bold = True
    p.add_run(f"{tokens:,} tokens")
    
    add_horizontal_line(doc)

    # Pages loop
    for page in json_data.get("pages", []):
        page_num = page.get("page_number")
        doc.add_heading(f"Page {page_num}", level=1)

        # Annotated image
        annotated_img = page.get("annotated_image", "")
        if annotated_img and os.path.exists(annotated_img):
            doc.add_paragraph("📄 Annotated Layout Preview:")
            doc.add_picture(annotated_img, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_horizontal_line(doc)

        # Elements loop
        for elem in page.get("elements", []):
            etype = elem.get("type")
            label = elem.get("label", "")
            conf  = elem.get("confidence", 0)

            if etype == "text":
                content = elem.get("content", "").strip()
                if content:
                    if label in ("header", "title"):
                        doc.add_heading(content, level=2)
                    else:
                        p = doc.add_paragraph(content)
                        if p.runs:
                            p.runs[0].font.size = Pt(10)

            elif etype == "table":
                add_label_badge(doc, f"TABLE | confidence: {conf:.2f} | {label}", color_hex="375623")
                img_path = elem.get("image_path", "")
                if img_path and os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(5.0))
                
                structured = elem.get("structured_data")
                if structured and (structured.get("rows") or structured.get("headers")):
                    method = structured.get("method", "unknown")
                    doc.add_paragraph(f"Structured Table Data (via {method}):")
                    add_table_to_doc(doc, structured)
                else:
                    doc.add_paragraph("⚠ No structured data could be extracted from this table.")
                add_horizontal_line(doc)

            elif etype == "image":
                add_label_badge(doc, f"IMAGE | confidence: {conf:.2f} | {label}", color_hex="7B2C2C")
                img_path = elem.get("image_path", "")
                if img_path and os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(4.5))
                
                # Vision description
                description = elem.get("vision_description")
                if description:
                    doc.add_heading("Vision Analysis (Moondream2):", level=3)
                    doc.add_paragraph(description).runs[0].font.size = Pt(10)
                add_horizontal_line(doc)

        doc.add_page_break()

    doc.save(output_path)
    return output_path
