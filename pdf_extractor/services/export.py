import os
import io
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def generate_word_report_stream(json_data):
    """Generates a simple, minimalist Word report with no fancy styling."""
    doc = Document()
    pdf_name = Path(json_data.get("pdf", "document")).stem
    
    doc.add_heading(f"Technical Extraction: {pdf_name}", level=0)

    # 1. Process Intelligence Summary (PER PDF METRICS)
    doc.add_heading("Process Intelligence Summary", level=1)
    
    proc_time = json_data.get("processing_time", 0)
    time_str = f"{proc_time:.1f}s" if proc_time < 60 else f"{int(proc_time//60)}m {int(proc_time%60)}s"

    doc.add_paragraph(f"• Total Processing Time: {time_str}")
    doc.add_paragraph(f"• Total Intelligence Content: {json_data.get('token_usage', 0):,} tokens")
    doc.add_paragraph(f"• Total AI Model Interactions: {json_data.get('request_count', 0)} requests")

    # 2. Sequential Data Pages
    for page in json_data.get("pages", []):
        page_num = page.get("page_number")
        doc.add_heading(f"Page {page_num}", level=1)

        for elem in page.get("elements", []):
            etype = elem.get("type")
            label = elem.get("label", "")

            if etype == "text":
                content = elem.get("content", "").strip()
                if content:
                    if label in ("header", "title", "section_title"):
                        doc.add_heading(content, level=2)
                    else:
                        doc.add_paragraph(content)

            elif etype == "table":
                doc.add_heading(f"Table: {label}", level=3)
                img_path = elem.get("image_path", "")
                if img_path and os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(4.0))
                
                s = elem.get("structured_data")
                if s and (s.get("rows") or s.get("headers")):
                    headers = s.get("headers", [])
                    rows = s.get("rows", [])
                    col_count = max(len(headers), max((len(r) for r in rows), default=0))
                    
                    if col_count > 0:
                        table = doc.add_table(rows=1 + len(rows), cols=col_count)
                        table.style = 'Table Grid'
                        
                        # Plain Header Row
                        if headers:
                            hdr_cells = table.rows[0].cells
                            for i, h in enumerate(headers):
                                if i < col_count: hdr_cells[i].text = str(h)
                        
                        # Plain Data Rows
                        for r_idx, row_data in enumerate(rows):
                            row_cells = table.rows[r_idx + 1].cells
                            for c_idx, val in enumerate(row_data):
                                if c_idx < col_count: row_cells[c_idx].text = str(val)

            elif etype == "image":
                desc = elem.get("vision_description", "")
                
                # SKIP: If the description is inconclusive (from my safety net), don't include it in the report.
                if "Summary Skipped/Inconclusive" in desc:
                    continue

                doc.add_heading(f"Visual Analysis: {label}", level=3)
                img_path = elem.get("image_path", "")
                
                if img_path and os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(3.2))
                
                if desc:
                    doc.add_paragraph(desc)

        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
